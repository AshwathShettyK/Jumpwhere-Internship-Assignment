from io import BytesIO
from docx import Document


class WordService:
    @staticmethod
    def render_docx(resume_data: dict) -> BytesIO:
        doc = Document()
        doc.add_heading(resume_data.get('name', ''), level=1)
        doc.add_paragraph(resume_data.get('designation', ''))
        doc.add_heading('Professional Summary', level=2)
        for b in resume_data.get('professional_summary', []):
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(b)
        doc.add_heading('Technical Skills', level=2)
        doc.add_paragraph(', '.join(resume_data.get('coding', []) + resume_data.get('tools', [])))
        doc.add_heading('Projects', level=2)
        for p in resume_data.get('projects', []):
            doc.add_heading(p.get('name', ''), level=3)
            techs = ', '.join(p.get('technologies', []) + p.get('tools', []))
            doc.add_paragraph(techs)
            doc.add_paragraph(p.get('description', ''))
            for r in p.get('role_responsibilities', []):
                pr = doc.add_paragraph(style='List Bullet')
                pr.add_run(r)
        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)
        return bio
