import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

class WordExportService:
    @staticmethod
    def generate_docx(context):
        document = docx.Document()
        
        # Configure page margins (0.8 inches on all sides)
        section = document.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
        # Configure default Normal style font to Calibri 11pt
        style = document.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        font.color.rgb = docx.shared.RGBColor(0x33, 0x33, 0x33) # Dark Charcoal
        
        employee = context['employee']
        
        # Add Header block: Name (Left-aligned) and Designation (Right-aligned) using a 2-column table
        header_table = document.add_table(rows=1, cols=2)
        header_table.autofit = False
        header_table.columns[0].width = Inches(4.5)
        header_table.columns[1].width = Inches(2.4)
        
        # Clean cell margins
        for row in header_table.rows:
            for cell in row.cells:
                cell.paragraphs[0].paragraph_format.space_after = Pt(0)
                cell.paragraphs[0].paragraph_format.space_before = Pt(0)
                
        # Cell 1: Employee Name
        cell_left = header_table.cell(0, 0)
        p_name = cell_left.paragraphs[0]
        run_name = p_name.add_run(employee.full_name)
        run_name.font.name = 'Calibri'
        run_name.font.size = Pt(18)
        run_name.font.bold = True
        run_name.font.color.rgb = docx.shared.RGBColor(0x0f, 0x17, 0x2a) # Dark slate
        
        # Cell 2: Designation
        cell_right = header_table.cell(0, 1)
        p_desig = cell_right.paragraphs[0]
        p_desig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        designation_name = employee.designation.designation_name if employee.designation else "Software Developer"
        run_desig = p_desig.add_run(f"Designation: {designation_name}")
        run_desig.font.name = 'Calibri'
        run_desig.font.size = Pt(11)
        run_desig.font.bold = True
        run_desig.font.color.rgb = docx.shared.RGBColor(0x47, 0x55, 0x69) # Slate gray
        
        # Add a horizontal line divider
        p_divider = document.add_paragraph()
        p_divider.paragraph_format.space_before = Pt(4)
        p_divider.paragraph_format.space_after = Pt(12)
        run_div = p_divider.add_run("_________________________________________________________________________________")
        run_div.font.color.rgb = docx.shared.RGBColor(0xcb, 0xd5, 0xe1) # light gray divider
        run_div.font.size = Pt(8)
        
        # Helper to add bold section headers
        def add_section_header(title):
            p = document.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(title)
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = docx.shared.RGBColor(0x0f, 0x17, 0x2a)
            return p
            
        # 1. Professional Summary Section
        add_section_header("Professional Summary:")
        if context['summary_bullets']:
            for bullet in context['summary_bullets']:
                p = document.add_paragraph(style='List Bullet')
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.space_before = Pt(0)
                run = p.add_run(bullet)
                run.font.size = Pt(11)
        else:
            p = document.add_paragraph("No professional summary bullet points configured.")
            p.paragraph_format.space_after = Pt(6)
            
        # 2. Technical Skill Set Section
        add_section_header("Technical Skill Set:")
        
        # Coding bullet
        p_coding = document.add_paragraph(style='List Bullet')
        p_coding.paragraph_format.space_after = Pt(3)
        p_coding.paragraph_format.space_before = Pt(0)
        run_coding_lbl = p_coding.add_run("Coding: ")
        run_coding_lbl.bold = True
        p_coding.add_run(context['coding_skills'])
        
        # Tools bullet
        p_tools = document.add_paragraph(style='List Bullet')
        p_tools.paragraph_format.space_after = Pt(12)
        p_tools.paragraph_format.space_before = Pt(0)
        run_tools_lbl = p_tools.add_run("Tools: ")
        run_tools_lbl.bold = True
        p_tools.add_run(context['tools'])

        # 3. Professional Projects Section
        add_section_header("Professional Projects")
        
        if context['projects']:
            for proj in context['projects']:
                # Project Header Line, e.g. "Project 1: Time tag"
                p_proj_head = document.add_paragraph()
                p_proj_head.paragraph_format.space_before = Pt(10)
                p_proj_head.paragraph_format.space_after = Pt(4)
                p_proj_head.paragraph_format.keep_with_next = True
                run_phead = p_proj_head.add_run(f"Project {proj['index']}: {proj['name']}")
                run_phead.font.bold = True
                run_phead.font.size = Pt(11)
                
                # Technology used bullet
                p_tech = document.add_paragraph(style='List Bullet')
                p_tech.paragraph_format.space_after = Pt(2)
                p_tech.paragraph_format.space_before = Pt(0)
                run_tech_lbl = p_tech.add_run("Technology used: ")
                run_tech_lbl.bold = True
                p_tech.add_run(proj['tech_used'])
                
                # Description bullet
                p_desc = document.add_paragraph(style='List Bullet')
                p_desc.paragraph_format.space_after = Pt(2)
                p_desc.paragraph_format.space_before = Pt(0)
                run_desc_lbl = p_desc.add_run("Description: ")
                run_desc_lbl.bold = True
                p_desc.add_run(proj['description'])
                
                # Role and Responsibilities bullet
                p_role = document.add_paragraph(style='List Bullet')
                p_role.paragraph_format.space_after = Pt(2)
                p_role.paragraph_format.space_before = Pt(0)
                run_role_lbl = p_role.add_run("Role and Responsibilities:")
                run_role_lbl.bold = True
                
                # Inner responsibilities sub-bullets
                if proj['responsibilities']:
                    for resp in proj['responsibilities']:
                        p_resp = document.add_paragraph(style='List Bullet')
                        p_resp.paragraph_format.space_after = Pt(2)
                        p_resp.paragraph_format.space_before = Pt(0)
                        # Setting left_indent creates a nested indent layout
                        p_resp.paragraph_format.left_indent = Inches(0.5)
                        p_resp.add_run(resp)
                else:
                    p_resp = document.add_paragraph(style='List Bullet')
                    p_resp.paragraph_format.space_after = Pt(2)
                    p_resp.paragraph_format.space_before = Pt(0)
                    p_resp.paragraph_format.left_indent = Inches(0.5)
                    p_resp.add_run(f"Worked as {proj['role']}.")
        else:
            p = document.add_paragraph("No project assignments available.")
            p.paragraph_format.space_after = Pt(6)
            
        return document
